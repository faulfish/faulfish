import openai
import os
import pandas as pd
import docx
import socket

def check_network():
    try:
        socket.create_connection(("www.google.com", 80))
        print("網路連接正常")
    except OSError:
        print("無法連接到網路")

check_network()

# 設置 OpenAI API 金鑰
api_key = os.getenv("OPENAI_API_KEY")

# 讀取 Excel 檔案
df = pd.read_excel("hot_keywords.xlsx", sheet_name=-1)

# 取得最後一個工作表的內容
last_sheet_name = list(df.keys())[-1]
# content = df.astype(str).values.tolist()
content = ["\n".join(row) for row in df.astype(str).values.tolist()]
# content = df["Keyword"].astype(str).tolist()
# print(content)

# 與 ChatGPT API 進行互動
response = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        # {"role": "user", "content": "Please provide a summary of the content related to the economy from below."},
        {"role": "user", "content": 
"Here are the suggested categories and translations for the provided news: entertainment, sports, politics, financial news, and other. Please remove items related to the entertainment, sports, and politics categories, and display the remaining translations along with the original text."},
        {"role": "assistant", "content": "\n".join(content)},
    ],
    max_tokens=300,
    n=1,
    stop=None,
    temperature=0.6,
    api_key=api_key
)

# 提取助理的回答作為摘要結果
summary = response.choices[0].message.content

# 列印摘要結果
print(summary)

# 創建 Word 文件
doc = docx.Document()

# 添加標題
doc.add_heading("Summary", level=1)

# 添加內容
doc.add_paragraph(summary)

# 保存 Word 檔案
doc.save("summary.docx")

